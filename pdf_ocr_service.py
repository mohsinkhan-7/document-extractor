"""PDF OCR & Chapter Extraction Service (Arabic)

Responsibilities:
 - Convert PDF pages to images (pdf2image + poppler)
 - Run Tesseract OCR (Arabic) on each page
 - Clean & normalize Arabic text
 - Detect chapter boundaries (الفصل / باب / chapter markers / numbered headings)
 - Export:
     * Word document (chapter-wise, cleaned, RTL paragraphs)
     * JSON structure of chapters
 - Environment diagnostics (presence of poppler & Tesseract Arabic traineddata)

Exceptions:
 - OCRConfigurationError: Misconfiguration / missing external binaries
 - OCRDependencyError: Missing Python package at runtime

NOTE: This is intentionally light-weight and avoids heavy NLP dependencies.
"""
from __future__ import annotations

import os
import re
import tempfile
from dataclasses import dataclass
import zipfile
import shutil
from typing import List, Dict, Any

try:
    from pdf2image import convert_from_path  # type: ignore
except Exception:  # pragma: no cover - handled in diagnostics
    convert_from_path = None  # type: ignore

try:
    import pytesseract  # type: ignore
except Exception:  # pragma: no cover
    pytesseract = None  # type: ignore

try:
    from PIL import Image  # type: ignore  # noqa: F401 - used implicitly by pdf2image
except Exception:  # pragma: no cover
    Image = None  # type: ignore

try:
    import arabic_reshaper  # type: ignore
    from bidi.algorithm import get_display  # type: ignore
except Exception:  # pragma: no cover
    arabic_reshaper = None  # type: ignore
    def get_display(txt: str) -> str:  # type: ignore
        return txt

try:
    from docx import Document  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
except Exception:  # pragma: no cover
    Document = None  # type: ignore


class OCRConfigurationError(RuntimeError):
    pass


class OCRDependencyError(RuntimeError):
    pass


POPPLER_CANDIDATES = [
    os.getenv("POPPLER_PATH"),
    r"C:\\poppler\\bin",
    r"C:\\poppler\\Library\\bin",
    r"C:\\tools\\poppler\\bin",
    r"C:\\tools\\poppler\\Library\\bin",
]


def _resolve_poppler_path() -> str | None:
    for p in POPPLER_CANDIDATES:
        if p and os.path.isdir(p) and os.path.exists(os.path.join(p, "pdfinfo.exe")):
            return p
    return None


def _configure_tesseract():
    if pytesseract is None:
        raise OCRDependencyError("pytesseract not installed. Check requirements.txt installation.")
    tess_env = os.getenv("TESSERACT_CMD")
    if tess_env:
        pytesseract.pytesseract.tesseract_cmd = tess_env


def diagnose_environment() -> Dict[str, Any]:
    """Return a JSON-friendly diagnostics snapshot used by /diagnostics/ocr endpoint."""
    poppler_path = _resolve_poppler_path()
    tess_cmd = None
    arabic_traineddata = False
    errors: List[str] = []

    if pytesseract:
        tess_cmd = getattr(pytesseract.pytesseract, "tesseract_cmd", "tesseract")
        # Try to infer tessdata location; this is heuristic.
        possible_dirs = [
            os.path.dirname(tess_cmd) if tess_cmd else "",
            os.path.join(os.path.dirname(tess_cmd) if tess_cmd else "", "tessdata"),
            os.getenv("TESSDATA_PREFIX", ""),
        ]
        for d in possible_dirs:
            if d and os.path.isdir(d):
                if os.path.exists(os.path.join(d, "ara.traineddata")):
                    arabic_traineddata = True
                    break
    else:
        errors.append("pytesseract import failed")

    if convert_from_path is None:
        errors.append("pdf2image import failed")
    if arabic_reshaper is None:
        errors.append("arabic_reshaper / python-bidi import failed (Arabic shaping issues likely)")
    if Document is None:
        errors.append("python-docx import failed (DOCX export disabled)")

    return {
        "poppler_path_detected": bool(poppler_path),
        "poppler_path": poppler_path,
        "tesseract_cmd": tess_cmd,
        "arabic_traineddata_found": arabic_traineddata,
        "errors": errors,
    }


_CHAPTER_PATTERNS = [
    re.compile(r"^(?:الفصل|باب)\s+([\u0621-\u064A0-9]+)"),  # الفصل / باب + word/number
    re.compile(r"^(?:فصل)\s*[:\-]?\s*([\u0621-\u064A0-9]+)?"),
    re.compile(r"^\s*[\d٠-٩]{1,3}\s*[-.،)]\s+"),  # numbered list item (Arabic/Latin digits)
]


ARABIC_SPACE_RE = re.compile(r"[ \t\u00A0]+")
MULTI_DOT_RE = re.compile(r"\.\.+")
TASHKEEL_RE = re.compile(r"[\u0610-\u061A\u064B-\u065F\u0670\u06D6-\u06ED]")  # optional diacritics removal


def _normalize_line(line: str, drop_diacritics: bool = False, aggressive: bool = False) -> str:
    line = line.strip('\ufeff\u200f\u200e')
    # unify spaces
    line = ARABIC_SPACE_RE.sub(" ", line)
    # collapse ellipsis
    line = MULTI_DOT_RE.sub("…", line)
    # normalize Arabic forms (simple subset) only when aggressive cleaning is requested
    if aggressive:
        replacements = {
            "ى": "ي",
            "إ": "ا",
            "أ": "ا",
            "ٱ": "ا",
            "ؤ": "و",
            "ئ": "ي",
        }
        for src, dst in replacements.items():
            line = line.replace(src, dst)
    if drop_diacritics:
        line = TASHKEEL_RE.sub("", line)
    return line.strip()


def _is_chapter_heading(line: str) -> bool:
    if not line:
        return False
    for pat in _CHAPTER_PATTERNS:
        if pat.search(line):
            return True
    # Heuristic: short line (<= 40 chars) with high proportion of Arabic letters & no period
    if len(line) <= 40 and ' ' in line and '۔' not in line and '.' not in line:
        arabic_letters = re.findall(r"[\u0621-\u064A]", line)
        if arabic_letters and len(arabic_letters) / max(len(line), 1) > 0.4:
            return True
    return False


@dataclass
class Chapter:
    title: str
    content: List[str]
    page_start: int
    page_end: int

    def as_dict(self) -> Dict[str, Any]:
        return {
            "title": self.title,
            "content": "\n".join(self.content).strip(),
            "page_start": self.page_start,
            "page_end": self.page_end,
        }


def _arabic_shape(text: str) -> str:
    if arabic_reshaper is None:
        return text
    try:
        # Avoid reshaping mixed content (emails/URLs/Latin-heavy) – preserve as-is
        if _looks_mixed_content(text):
            return text
        reshaped = arabic_reshaper.reshape(text)
        return get_display(reshaped)
    except Exception:  # pragma: no cover
        return text


def _looks_mixed_content(text: str) -> bool:
    has_ar = re.search(r"[\u0621-\u064A]", text) is not None
    has_lat = re.search(r"[A-Za-z]", text) is not None
    has_email_or_url = ("@" in text) or ("http://" in text) or ("https://" in text) or ("www." in text)
    return (has_ar and has_lat) or has_email_or_url


def _ocr_pdf_to_pages(pdf_path: str, dpi: int = 250, lang: str | None = None) -> List[str]:
    poppler_path = _resolve_poppler_path()
    if convert_from_path is None:
        raise OCRDependencyError("pdf2image not available. Install dependencies.")
    if pytesseract is None:
        raise OCRDependencyError("pytesseract not available. Install dependencies.")

    _configure_tesseract()
    if lang is None:
        lang = os.getenv("OCR_LANG", "ara+eng")

    try:
        # If Poppler not found on Windows give clearer guidance before attempting
        if os.name == 'nt' and not poppler_path:
            raise OCRConfigurationError(
                "Poppler (pdfinfo.exe / pdftoppm.exe) not found. Install from https://github.com/oschwartz10612/poppler-windows/releases, "
                "extract e.g. to C\\\poppler, then set POPPLER_PATH to the 'bin' folder (temporary PowerShell: $env:POPPLER_PATH='C:\\poppler\\bin'). "
                "After setting, restart the server or shell."
            )
        images = convert_from_path(pdf_path, dpi=dpi, poppler_path=poppler_path)
    except Exception as e:  # pragma: no cover - environment specific
        raise OCRConfigurationError(f"Failed converting PDF pages: {e}")

    texts: List[str] = []
    for img in images:
        try:
            txt = pytesseract.image_to_string(img, lang=lang)
        except Exception as e:  # pragma: no cover
            raise OCRConfigurationError(f"Tesseract OCR failed: {e}")
        # Basic cleanup
        # Keep normalization light to preserve mixed content fidelity
        tmp_lines = []
        for l in txt.splitlines():
            nl = _normalize_line(l, drop_diacritics=False, aggressive=False)
            if nl != "":
                tmp_lines.append(nl)
        lines = tmp_lines
        texts.append("\n".join(lines))
    return texts


def _segment_chapters(pages: List[str], page_offset: int = 0) -> List[Chapter]:
    chapters: List[Chapter] = []
    current = Chapter(title="مقدمة", content=[], page_start=1 + page_offset, page_end=1 + page_offset)
    for idx, page_text in enumerate(pages, start=1):
        lines = page_text.splitlines()
        for line in lines:
            if _is_chapter_heading(line) and current.content:
                # finalize current
                current.page_end = idx + page_offset
                chapters.append(current)
                current = Chapter(title=line, content=[], page_start=idx + page_offset, page_end=idx + page_offset)
            else:
                current.content.append(line)
        current.page_end = idx + page_offset
    if current.content:
        chapters.append(current)
    return chapters


def extract_chapters_as_json(pdf_path: str, start_page: int = 1) -> List[Dict[str, Any]]:
    pages = _ocr_pdf_to_pages(pdf_path)
    page_offset = 0
    if start_page > 1:
        page_offset = start_page - 1
        pages = pages[page_offset:]
    chapters = _segment_chapters(pages, page_offset=page_offset)
    shaped = []
    for ch in chapters:
        shaped.append({
            "title": _arabic_shape(ch.title),
            "content": _arabic_shape(ch.as_dict()["content"]),
            "page_start": ch.page_start,
            "page_end": ch.page_end,
        })
    return shaped


def _set_paragraph_rtl(paragraph):  # type: ignore
    try:
        p = paragraph._p  # noqa: SLF001
        pPr = p.get_or_add_pPr()
        bidi = OxmlElement('w:bidi')
        pPr.append(bidi)
    except Exception:  # pragma: no cover
        pass


def pdf_to_word_chapters(pdf_path: str, output_docx: str, start_page: int = 1) -> str:
    if Document is None:
        raise OCRDependencyError("python-docx not installed.")
    pages = _ocr_pdf_to_pages(pdf_path)
    page_offset = 0
    if start_page > 1:
        page_offset = start_page - 1
        pages = pages[page_offset:]
    chapters = _segment_chapters(pages, page_offset=page_offset)
    doc = Document()
    doc.core_properties.title = os.path.basename(pdf_path)
    shape_for_docx = os.getenv("DOCX_SHAPE_ARABIC", "false").lower() == "true"
    for ch in chapters:
        heading_text = _arabic_shape(ch.title) if (shape_for_docx and not _looks_mixed_content(ch.title)) else ch.title
        h = doc.add_heading(heading_text, level=1)
        _set_paragraph_rtl(h)
        paragraph_blocks = "\n".join(ch.content).split("\n\n")
        for block in paragraph_blocks:
            txt = block.strip()
            if not txt:
                continue
            out_txt = _arabic_shape(txt) if (shape_for_docx and not _looks_mixed_content(txt)) else txt
            p = doc.add_paragraph(out_txt)
            _set_paragraph_rtl(p)
    doc.save(output_docx)
    return output_docx


def chapters_json_to_word(chapters: List[Dict[str, Any]], output_docx: str) -> str:
    """Create a DOCX from a chapters JSON (title, content)."""
    if Document is None:
        raise OCRDependencyError("python-docx not installed.")
    if not chapters:
        raise ValueError("No chapters provided")
    doc = Document()
    doc.core_properties.title = "Chapters Export"
    shape_for_docx = os.getenv("DOCX_SHAPE_ARABIC", "false").lower() == "true"
    for ch in chapters:
        raw_title = str(ch.get("title", ""))
        title = (_arabic_shape(raw_title) if (shape_for_docx and not _looks_mixed_content(raw_title)) else raw_title) or "(بدون عنوان)"
        content = str(ch.get("content", "")).strip()
        h = doc.add_heading(title, level=1)
        _set_paragraph_rtl(h)
        # Split paragraphs by blank lines
        for para in re.split(r"\n{2,}", content):
            para = para.strip()
            if not para:
                continue
            out_txt = _arabic_shape(para) if (shape_for_docx and not _looks_mixed_content(para)) else para
            p = doc.add_paragraph(out_txt)
            _set_paragraph_rtl(p)
    doc.save(output_docx)
    return output_docx


def _sanitize_filename(name: str, fallback: str = "chapter") -> str:
    # Allow Arabic letters, Latin letters, digits, basic punctuation, replace spaces with underscores
    name = name.strip()
    name = re.sub(r"[\s\u200f\u200e]+", "_", name)
    name = re.sub(r"[^\u0621-\u064Aa-zA-Z0-9_\-()\[\]{}]+", "", name)
    name = name.strip("._- ")
    if not name:
        name = fallback
    return name[:120]


def chapter_to_word(chapter: Dict[str, Any], output_docx: str) -> str:
    if Document is None:
        raise OCRDependencyError("python-docx not installed.")
    doc = Document()
    doc.core_properties.title = str(chapter.get("title", "Chapter"))
    shape_for_docx = os.getenv("DOCX_SHAPE_ARABIC", "false").lower() == "true"
    raw_title = str(chapter.get("title", ""))
    title = (_arabic_shape(raw_title) if (shape_for_docx and not _looks_mixed_content(raw_title)) else raw_title) or "(بدون عنوان)"
    h = doc.add_heading(title, level=1)
    _set_paragraph_rtl(h)
    content = str(chapter.get("content", ""))
    for para in re.split(r"\n{2,}", content):
        para = para.strip()
        if not para:
            continue
        out_txt = _arabic_shape(para) if (shape_for_docx and not _looks_mixed_content(para)) else para
        p = doc.add_paragraph(out_txt)
        _set_paragraph_rtl(p)
    doc.save(output_docx)
    return output_docx


def export_chapters_to_zip(pdf_path: str, zip_path: str, start_page: int = 1) -> Dict[str, Any]:
    """Extract chapters from PDF and export each as a separate DOCX inside a ZIP.
    Returns metadata with list of files created.
    """
    chapters = extract_chapters_as_json(pdf_path, start_page=start_page)
    if not chapters:
        raise RuntimeError("No chapters detected")
    temp_dir = tempfile.mkdtemp(prefix="chapters_")
    created_files: List[str] = []
    used_names = set()
    try:
        for idx, ch in enumerate(chapters, start=1):
            base = _sanitize_filename(ch.get("title") or f"chapter_{idx}", fallback=f"chapter_{idx}")
            # Ensure unique
            name = base
            k = 1
            while name.lower() in used_names:
                name = f"{base}_{k}"
                k += 1
            used_names.add(name.lower())
            out_path = os.path.join(temp_dir, f"{idx:02d}_{name}.docx")
            chapter_to_word(ch, out_path)
            created_files.append(out_path)

        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for f in created_files:
                arcname = os.path.basename(f)
                zf.write(f, arcname)
    finally:
        # Keep temp_dir for now if needed; we could clean but then files aren’t accessible.
        # We'll clean temp_dir after zipping.
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass
    return {"zip_path": zip_path, "count": len(created_files)}


# --- TOC-based chapter export -------------------------------------------------

_TRAILING_PAGE_RE = re.compile(
    r"^(?P<title>.+?)\s*[\.\·•\-–—،…\s]*\s(?P<page>[0-9\u0660-\u0669]{1,4})$"
)

def _arabic_digits_to_int(s: str) -> int | None:
    # Convert Arabic-Indic digits to Latin then parse
    trans = str.maketrans("٠١٢٣٤٥٦٧٨٩", "0123456789")
    s2 = s.translate(trans)
    m = re.search(r"(\d+)", s2)
    if not m:
        return None
    try:
        return int(m.group(1))
    except Exception:
        return None
    
def detect_toc_page(pdf_path: str, max_scan_pages: int = 10, lang: str | None = None) -> int | None:
    """Scan first pages of PDF and try to detect TOC page by keywords."""
    poppler_path = _resolve_poppler_path()
    if convert_from_path is None:
        raise OCRDependencyError("pdf2image not available. Install dependencies.")
    if pytesseract is None:
        raise OCRDependencyError("pytesseract not available. Install dependencies.")

    _configure_tesseract()
    if lang is None:
        lang = os.getenv("OCR_LANG", "ara+eng")

    try:
        images = convert_from_path(
            pdf_path, first_page=1, last_page=max_scan_pages,
            poppler_path=poppler_path, dpi=150
        )
    except Exception as e:
        raise OCRConfigurationError(f"Failed scanning for TOC: {e}")

    keywords = ["المحتويات", "فهرس", "جدول المحتويات"]
    for idx, img in enumerate(images, start=1):
        try:
            txt = pytesseract.image_to_string(img, lang=lang)
        except Exception:
            continue
        for kw in keywords:
            if kw in txt:
                return idx
    return None


def extract_toc_entries(pdf_path: str, toc_page: int = 5, lang: str | None = None) -> List[Dict[str, Any]]:
    poppler_path = _resolve_poppler_path()
    if convert_from_path is None:
        raise OCRDependencyError("pdf2image not available. Install dependencies.")
    if pytesseract is None:
        raise OCRDependencyError("pytesseract not available. Install dependencies.")

    _configure_tesseract()
    if lang is None:
        lang = os.getenv("OCR_LANG", "ara+eng")

    try:
        # Use higher DPI for better OCR of Arabic fonts
        images = convert_from_path(
            pdf_path, first_page=toc_page, last_page=toc_page,
            poppler_path=poppler_path, dpi=300
        )
    except Exception as e:
        raise OCRConfigurationError(f"Failed reading TOC page: {e}")

    if not images:
        return []

    try:
        txt = pytesseract.image_to_string(images[0], lang=lang)
    except Exception as e:
        raise OCRConfigurationError(f"Tesseract OCR failed on TOC page: {e}")

    entries: List[Dict[str, Any]] = []
    for raw in txt.splitlines():
        line = _normalize_line(raw)
        if not line:
            continue

        # Collapse repeating dots/dashes
        line = re.sub(r"[\.·•،…\-–—]{2,}", " … ", line)

        m = _TRAILING_PAGE_RE.search(line)
        if not m:
            continue

        title = m.group("title").strip(". \t").strip()
        page_s = m.group("page")
        page_i = _arabic_digits_to_int(page_s)
        if page_i is None:
            continue

        entries.append({
            "title": title,
            "printed_page": page_i
        })

    return entries



def export_chapters_to_zip_from_toc(pdf_path: str, zip_path: str, toc_page: int = 5, printed_to_pdf_offset: int = 0) -> Dict[str, Any]:
    # OCR all pages once; slice per chapter using TOC page numbers
    
    pages = _ocr_pdf_to_pages(pdf_path)
    # toc = extract_toc_entries(pdf_path, toc_page=toc_page)
    # Auto-detect TOC if not provided explicitly
    if toc_page is None or toc_page <= 0:
        detected = detect_toc_page(pdf_path)
        if detected:
            toc_page = detected

    # OCR all pages once; slice per chapter using TOC page numbers
    pages = _ocr_pdf_to_pages(pdf_path)
    toc = extract_toc_entries(pdf_path, toc_page=toc_page)

    if not toc:
        # Fallback: return the raw OCR text for debugging instead of 500
        try:
            poppler_path = _resolve_poppler_path()
            images = convert_from_path(pdf_path, first_page=toc_page, last_page=toc_page, poppler_path=poppler_path, dpi=300)
            raw_txt = pytesseract.image_to_string(images[0], lang=os.getenv("OCR_LANG", "ara+eng"))
        except Exception as e:
            raw_txt = f"[OCR failed: {e}]"

        return {
            "zip_path": None,
            "count": 0,
            "toc_count": 0,
            "raw_toc_text": raw_txt
        }

    # Compute start indices in 1-based PDF page numbers, then to 0-based for list slicing
    starts: List[Dict[str, Any]] = []
    for e in toc:
        pdf_page_num = e["printed_page"] + printed_to_pdf_offset
        if pdf_page_num < 1 or pdf_page_num > len(pages):
            continue
        starts.append({"title": e["title"], "pdf_page": pdf_page_num})
    # Ensure sorted by page
    starts.sort(key=lambda x: x["pdf_page"])
    if not starts:
        raise RuntimeError("TOC pages map outside the PDF range. Adjust the offset.")

    temp_dir = tempfile.mkdtemp(prefix="chapters_toc_")
    created_files: List[str] = []
    used_names = set()
    try:
        for i, s in enumerate(starts):
            start_p = s["pdf_page"]
            end_p = (starts[i + 1]["pdf_page"] - 1) if i + 1 < len(starts) else len(pages)
            content = "\n".join(pages[start_p - 1:end_p])
            ch = {"title": s["title"], "content": content, "page_start": start_p, "page_end": end_p}
            base = _sanitize_filename(ch["title"] or f"chapter_{i+1}", fallback=f"chapter_{i+1}")
            name = base
            k = 1
            while name.lower() in used_names:
                name = f"{base}_{k}"
                k += 1
            used_names.add(name.lower())
            out_path = os.path.join(temp_dir, f"{i+1:02d}_{name}.docx")
            chapter_to_word(ch, out_path)
            created_files.append(out_path)

        with zipfile.ZipFile(zip_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for f in created_files:
                zf.write(f, os.path.basename(f))
    finally:
        try:
            shutil.rmtree(temp_dir, ignore_errors=True)
        except Exception:
            pass
    return {"zip_path": zip_path, "count": len(created_files), "toc_count": len(toc)}


def main_cli():  # pragma: no cover - convenience
    import argparse
    parser = argparse.ArgumentParser(description="Arabic PDF OCR to Word (chapter-wise)")
    parser.add_argument("pdf", help="Input PDF path")
    parser.add_argument("--out", default=None, help="Output DOCX path")
    parser.add_argument("--json", action="store_true", help="Print chapters JSON and exit")
    parser.add_argument("--dpi", type=int, default=250)
    args = parser.parse_args()

    if not os.path.exists(args.pdf):
        raise SystemExit("Input PDF not found")
    if args.json:
        data = extract_chapters_as_json(args.pdf)
        import json
        print(json.dumps(data, ensure_ascii=False, indent=2))
        return
    out = args.out or os.path.splitext(args.pdf)[0] + "_chapters.docx"
    pdf_to_word_chapters(args.pdf, out)
    print(f"Written: {out}")


if __name__ == "__main__":  # pragma: no cover
    main_cli()
